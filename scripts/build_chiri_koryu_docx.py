"""地理交流広場第10号 北木島訪問記の Word ファイル(.docx)を生成する。

テンプレート（コラム体裁）のセクションプロパティ・フォント設定をベースに、
draft.md の内容を python-docx で組み直す。

実行方法:
    uv run python scripts/build_chiri_koryu_docx.py

入力:
    docs/articles/2026_chiri-koryu-10/draft.md
    /Users/otsuka/Downloads/地理交流広場テンプレート（コラム体裁）_v3.docx

出力:
    docs/articles/2026_chiri-koryu-10/draft.docx (.gitignore で除外)

体裁:
    - B5 (182×257mm)
    - 余白: 上下25mm、左右18mm
    - 本文: UD デジタル 教科書体 N-R 9pt
    - 見出し: UD デジタル 教科書体 N-B
    - タイトル: UD デジタル 教科書体 N-B 12pt
"""

from __future__ import annotations

import re
import shutil
from copy import deepcopy
from pathlib import Path
from typing import Iterator

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Mm, Pt
from docx.oxml import OxmlElement

PROJECT_ROOT = Path(__file__).resolve().parents[1]
ARTICLE_DIR = PROJECT_ROOT / "docs" / "articles" / "2026_chiri-koryu-10"
DRAFT_MD = ARTICLE_DIR / "draft.md"
OUT_DOCX = ARTICLE_DIR / "draft.docx"
TEMPLATE = Path("/Users/otsuka/Downloads/地理交流広場テンプレート（コラム体裁）_v3.docx")

FONT_BOLD = "UD デジタル 教科書体 N-B"
FONT_REGULAR = "UD デジタル 教科書体 N-R"


def set_run_font(run, name: str, size_pt: float, bold: bool = False) -> None:
    """ランに UD デジタル教科書体 を直接指定する（ASCII・eastAsia 両方）。"""
    run.font.name = name
    rpr = run._element.get_or_add_rPr()
    # 既存の rFonts を削除して新しい設定を入れる
    for elem in rpr.findall(qn("w:rFonts")):
        rpr.remove(elem)
    rfonts = OxmlElement("w:rFonts")
    rfonts.set(qn("w:ascii"), name)
    rfonts.set(qn("w:eastAsia"), name)
    rfonts.set(qn("w:hAnsi"), name)
    rfonts.set(qn("w:cs"), name)
    rfonts.set(qn("w:hint"), "eastAsia")
    rpr.insert(0, rfonts)
    run.font.size = Pt(size_pt)
    if bold:
        run.bold = True


def _add_clear_wrap_spacer(doc) -> None:
    """空の段落に clear="all" のテキスト改行を入れて、フロート画像の下から始める。

    本段落と見出しを分離することで、見出し本文に余分な break が混ざらないようにし、
    日本語見出しの文字幅膨張を回避する。
    """
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_before = Pt(0)
    spacer.paragraph_format.space_after = Pt(0)
    sp_run = spacer.add_run()
    br = OxmlElement("w:br")
    br.set(qn("w:type"), "textWrapping")
    br.set(qn("w:clear"), "all")
    sp_run._element.append(br)


def add_paragraph(doc, text: str, *, font: str = FONT_REGULAR, size_pt: float = 9.0,
                  bold: bool = False, align: str | None = None,
                  space_before: float = 0.0, space_after: float = 0.0,
                  first_line_indent: bool = False,
                  clear_wrap: bool = False,
                  snap_to_grid: bool | None = None):
    """本文段落を追加する。インライン記法（**bold**, ^sup^）も簡易処理する。

    clear_wrap=True にすると、段落の **前に** 別途 clear="all" の空段落を挿入し、
    上にある回り込み画像の下から本段落が始まるようにする（見出し等で使用）。
    本段落自体には break を入れないので、文字幅は自然な詰まり方になる。

    snap_to_grid=False を指定すると、文書グリッドへの文字スナップを無効化する。
    本文と異なるサイズの見出し（11pt/9.5pt 等）でグリッド幅に文字が引き伸ばされる
    現象を防ぐため、見出し段落で False を指定する。
    """
    if clear_wrap:
        _add_clear_wrap_spacer(doc)

    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    if first_line_indent:
        pf.first_line_indent = Pt(size_pt)  # 1文字分の字下げ
    if align == "center":
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif align == "right":
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    elif align == "left":
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # snapToGrid 制御（見出しなど、本文と異なるサイズの段落で off にする）
    if snap_to_grid is False:
        ppr = p._element.get_or_add_pPr()
        snap = OxmlElement("w:snapToGrid")
        snap.set(qn("w:val"), "0")
        ppr.append(snap)

    # インライン記法を解釈してランを作る
    for piece, kind in _parse_inline(text):
        run = p.add_run(piece)
        is_bold = bold or kind == "bold"
        set_run_font(run, font, size_pt, bold=is_bold)
        if kind == "sup":
            run.font.superscript = True
    return p


def _parse_inline(text: str) -> Iterator[tuple[str, str]]:
    """**bold**, ^sup^ を解釈してテキスト断片を返す。"""
    # トークン化: 順番に **bold** / ^sup^ / 通常テキスト
    pattern = re.compile(r"(\*\*[^*]+\*\*|\^[^^]+\^)")
    pos = 0
    for m in pattern.finditer(text):
        if m.start() > pos:
            yield text[pos:m.start()], "normal"
        token = m.group(0)
        if token.startswith("**"):
            yield token[2:-2], "bold"
        elif token.startswith("^"):
            yield token[1:-1], "sup"
        pos = m.end()
    if pos < len(text):
        yield text[pos:], "normal"


def _set_no_borders(tbl_pr) -> None:
    """テーブルの境界線をすべて非表示にする。"""
    borders = OxmlElement("w:tblBorders")
    for b in ("top", "left", "bottom", "right", "insideH", "insideV"):
        e = OxmlElement(f"w:{b}")
        e.set(qn("w:val"), "nil")
        borders.append(e)
    tbl_pr.append(borders)


def _make_table_floating(table, *, side: str = "right",
                         left_from_text: int = 180,
                         right_from_text: int = 180) -> None:
    """テーブル全体を回り込みフロートにする（画像+キャプションを一体で浮かべる）。

    side: "right" / "left"
    left_from_text/right_from_text は本文との余白（twip 単位、180=約3pt=1mm弱）
    """
    tbl_pr = table._element.find(qn("w:tblPr"))
    if tbl_pr is None:
        tbl_pr = OxmlElement("w:tblPr")
        table._element.insert(0, tbl_pr)

    tblpPr = OxmlElement("w:tblpPr")
    tblpPr.set(qn("w:leftFromText"), str(left_from_text))
    tblpPr.set(qn("w:rightFromText"), str(right_from_text))
    tblpPr.set(qn("w:topFromText"), "100")
    tblpPr.set(qn("w:bottomFromText"), "100")
    tblpPr.set(qn("w:vertAnchor"), "text")
    tblpPr.set(qn("w:horzAnchor"), "margin")
    tblpPr.set(qn("w:tblpXSpec"), side)
    tbl_pr.append(tblpPr)

    overlap = OxmlElement("w:tblOverlap")
    overlap.set(qn("w:val"), "never")
    tbl_pr.append(overlap)


def _add_picture_to_cell(cell, image_path: Path, width_mm: float) -> None:
    """セル内に画像を中央揃えで挿入。"""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf = p.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    run = p.add_run()
    run.add_picture(str(image_path), width=Mm(width_mm))


def _add_caption_to_cell(cell, caption: str) -> None:
    """セル内にキャプションを中央揃えで追加。"""
    p = cell.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf = p.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    run = p.add_run(caption)
    set_run_font(run, FONT_REGULAR, 9.0)


def add_image_floating(doc, image_path: Path, caption: str, *, side: str = "right",
                       width_mm: float = 70.0) -> None:
    """画像とキャプションを 1×1 テーブルにまとめて、回り込みフロート配置する。

    テンプレートの画像配置と同様、本文が画像（とキャプション）の反対側に回り込む。
    キャプションは画像の直下、テーブル内に確実に配置される。
    """
    table = doc.add_table(rows=1, cols=1)
    table.autofit = False
    cell = table.cell(0, 0)
    cell.width = Mm(width_mm + 4)

    _add_picture_to_cell(cell, image_path, width_mm)
    _add_caption_to_cell(cell, caption)

    # テーブルの幅を明示
    tbl_pr = table._element.find(qn("w:tblPr"))
    tbl_w = OxmlElement("w:tblW")
    tbl_w.set(qn("w:w"), str(int((width_mm + 4) * 56.7)))  # mm → twip
    tbl_w.set(qn("w:type"), "dxa")
    tbl_pr.append(tbl_w)

    _set_no_borders(tbl_pr)
    _make_table_floating(table, side=side)


def add_image_grid(doc, image_blocks: list[dict], *, width_mm: float = 65.0) -> None:
    """連続する複数の画像を 1×N テーブル（横並び）で配置する。

    回り込みではなく中央寄せのインラインテーブルとし、各セルに画像+キャプションを入れる。
    これによって連続画像のオーバーラップを防ぎ、左右に並べて表示する。
    """
    n = len(image_blocks)
    table = doc.add_table(rows=1, cols=n)
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, blk in enumerate(image_blocks):
        cell = table.cell(0, i)
        cell.width = Mm(width_mm + 4)
        img_path = ARTICLE_DIR / blk["path"]
        if img_path.exists():
            _add_picture_to_cell(cell, img_path, width_mm)
            _add_caption_to_cell(cell, blk["caption"])

    # 各セルの幅と境界線設定
    tbl_pr = table._element.find(qn("w:tblPr"))
    if tbl_pr is None:
        tbl_pr = OxmlElement("w:tblPr")
        table._element.insert(0, tbl_pr)
    tbl_w = OxmlElement("w:tblW")
    tbl_w.set(qn("w:w"), str(int((width_mm + 4) * n * 56.7)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_pr.append(tbl_w)
    _set_no_borders(tbl_pr)


def add_image_inline_full(doc, image_path: Path, caption: str,
                          width_mm: float = 140.0) -> None:
    """画像を中央寄せの inline 配置で大きく表示する（フル幅用）。

    位置関係を正確に把握する必要がある図（位置図、三スケール比較等）に使用。
    本文の回り込みは発生せず、図と次の段落は上下に配置される。
    """
    p_img = doc.add_paragraph()
    p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf = p_img.paragraph_format
    pf.space_before = Pt(6)
    pf.space_after = Pt(2)
    run = p_img.add_run()
    run.add_picture(str(image_path), width=Mm(width_mm))

    p_cap = doc.add_paragraph()
    p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cf = p_cap.paragraph_format
    cf.space_before = Pt(0)
    cf.space_after = Pt(8)
    cap_run = p_cap.add_run(caption)
    set_run_font(cap_run, FONT_REGULAR, 9.0)


def add_image(doc, image_path: Path, caption: str) -> None:
    """単独画像を回り込みフロートで配置（後方互換のため残す）。"""
    add_image_floating(doc, image_path, caption, side="right", width_mm=70.0)


# 横幅いっぱいで表示する図（位置関係や全体把握が重要なもの、または横長の図）
FULLWIDTH_FIGURES = {
    "fig02_location_map.png",
    "fig07_ndwi_histogram.png",
    "fig08_water_distribution.png",
    "fig09_multiscale.png",
}


def _is_fullwidth(blk: dict) -> bool:
    """そのブロックがフル幅 inline 表示にすべき画像か判定する。"""
    return Path(blk["path"]).name in FULLWIDTH_FIGURES


def clear_body_keep_sectPr(doc) -> None:
    """body の中身を全て削除する（最後の sectPr は document 側で保持される）。"""
    body = doc.element.body
    # sectPr 以外の全要素を削除
    sectPr = body.find(qn("w:sectPr"))
    for child in list(body):
        if child.tag != qn("w:sectPr"):
            body.remove(child)
    # sectPr は最後に再追加される必要があるが、find で取得した要素はまだ body にある
    # （find は remove しないので、上記 list(body) で sectPr も含まれていた場合は除外済）
    # 念のため: sectPr が残っているか確認、なければ追加
    if sectPr is not None and body.find(qn("w:sectPr")) is None:
        body.append(sectPr)


def parse_markdown(md_text: str) -> list[dict]:
    """draft.md を簡易パースしてブロックリストに変換する。"""
    blocks: list[dict] = []
    lines = md_text.split("\n")
    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # 空行はスキップ（次のループで判定）
        if stripped == "":
            i += 1
            continue

        # 区切り線
        if stripped == "---":
            blocks.append({"type": "hr"})
            i += 1
            continue

        # H1 タイトル
        if stripped.startswith("# "):
            blocks.append({"type": "h1", "text": stripped[2:]})
            i += 1
            continue

        # H2 章
        if stripped.startswith("## "):
            blocks.append({"type": "h2", "text": stripped[3:]})
            i += 1
            continue

        # H3 節
        if stripped.startswith("### "):
            blocks.append({"type": "h3", "text": stripped[4:]})
            i += 1
            continue

        # 画像 ![caption](path)
        m = re.match(r"!\[([^\]]+)\]\(([^)]+)\)", stripped)
        if m:
            blocks.append({"type": "image", "caption": m.group(1), "path": m.group(2)})
            i += 1
            continue

        # 文献リスト（- で始まる）
        if stripped.startswith("- "):
            blocks.append({"type": "list_item", "text": stripped[2:]})
            i += 1
            continue

        # 注（数字+）で始まる）
        m_note = re.match(r"^(\d+)\)\s+(.+)$", stripped)
        if m_note:
            blocks.append({"type": "note_item", "num": m_note.group(1), "text": m_note.group(2)})
            i += 1
            continue

        # 副題（―...―）
        if stripped.startswith("―") and stripped.endswith("―"):
            blocks.append({"type": "subtitle", "text": stripped})
            i += 1
            continue

        # それ以外は本文段落（連続する非空行を結合）
        para_lines = [stripped]
        j = i + 1
        while j < len(lines) and lines[j].strip() and not _is_block_start(lines[j].strip()):
            para_lines.append(lines[j].strip())
            j += 1
        blocks.append({"type": "para", "text": " ".join(para_lines)})
        i = j

    return blocks


def _is_block_start(line: str) -> bool:
    """この行が新しいブロックの開始か判定する。"""
    if line.startswith(("# ", "## ", "### ", "- ", "!")):
        return True
    if line == "---":
        return True
    if re.match(r"^\d+\)\s+", line):
        return True
    if line.startswith("―") and line.endswith("―"):
        return True
    return False


def is_author_line(text: str) -> bool:
    """著者行（短く、特定パターンに合う）を判定する。"""
    return text.strip() in ("Otsuka Noboru", "大塚　昇", "大塚 昇")


def render(blocks: list[dict], doc) -> None:
    """ブロックを Word ドキュメントに書き込む。

    連続する画像ブロックはまとめて検出し、1枚なら回り込みフロート、
    2枚以上なら横並びテーブル（インライン中央寄せ）として配置する。
    """
    # まず空のドキュメントから始まる前提。最初に表紙的に title/subtitle/author を扱う
    seen_h1 = False
    i = 0
    while i < len(blocks):
        blk = blocks[i]
        # 連続する画像はグループ化（ただしフル幅画像は単独扱いにする）
        if blk["type"] == "image":
            group = [blk]
            j = i + 1
            if not _is_fullwidth(blk):
                while (j < len(blocks) and blocks[j]["type"] == "image"
                       and not _is_fullwidth(blocks[j])):
                    group.append(blocks[j])
                    j += 1
            if len(group) == 1:
                img_path = ARTICLE_DIR / group[0]["path"]
                if img_path.exists():
                    if _is_fullwidth(group[0]):
                        # フル幅 inline 配置（位置図・三スケール比較）
                        add_image_inline_full(doc, img_path, group[0]["caption"],
                                              width_mm=140.0)
                    else:
                        # 通常は右寄せ float
                        add_image_floating(doc, img_path, group[0]["caption"],
                                           side="right", width_mm=70.0)
                else:
                    add_paragraph(doc, f"[画像が見つかりません: {group[0]['path']}]",
                                  font=FONT_REGULAR, size_pt=9.0)
            else:
                add_image_grid(doc, group, width_mm=65.0)
            i = j
            continue

        t = blk["type"]
        if t == "h1":
            # タイトル: UD N-B 12pt 中央揃え
            add_paragraph(doc, blk["text"], font=FONT_BOLD, size_pt=12.0,
                          align="center", space_before=0, space_after=4)
            seen_h1 = True
        elif t == "subtitle":
            # 副題: UD N-B 9pt 中央揃え
            add_paragraph(doc, blk["text"], font=FONT_BOLD, size_pt=9.0,
                          align="center", space_after=4)
        elif t == "para" and is_author_line(blk["text"]):
            # 著者行: 10.5pt 中央揃え
            add_paragraph(doc, blk["text"], font=FONT_REGULAR, size_pt=10.5,
                          align="center", space_after=12)
        elif t == "hr":
            # 区切り線はスキップ
            i += 1
            continue
        elif t == "h2":
            # 章: UD N-B 11pt 太字、左揃え、グリッドスナップ off（文字膨張防止）
            add_paragraph(doc, blk["text"], font=FONT_BOLD, size_pt=11.0,
                          bold=True, space_before=10, space_after=4,
                          align="left", snap_to_grid=False)
        elif t == "h3":
            # 節: UD N-B 9.5pt 太字、左揃え、グリッドスナップ off
            add_paragraph(doc, blk["text"], font=FONT_BOLD, size_pt=9.5,
                          bold=True, space_before=6, space_after=2,
                          align="left", snap_to_grid=False)
        elif t == "list_item":
            # 文献リスト
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Mm(5)
            p.paragraph_format.first_line_indent = Mm(-5)
            p.paragraph_format.space_after = Pt(2)
            for piece, kind in _parse_inline("・ " + blk["text"]):
                r = p.add_run(piece)
                set_run_font(r, FONT_REGULAR, 8.5, bold=(kind == "bold"))
        elif t == "note_item":
            # 注: 番号と本文
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Mm(5)
            p.paragraph_format.first_line_indent = Mm(-5)
            p.paragraph_format.space_after = Pt(2)
            num_run = p.add_run(f"{blk['num']}）")
            set_run_font(num_run, FONT_REGULAR, 8.0)
            for piece, kind in _parse_inline(blk["text"]):
                r = p.add_run(piece)
                set_run_font(r, FONT_REGULAR, 8.0, bold=(kind == "bold"))
        elif t == "para":
            # 本文段落 — 章タイトル「注」「文献」の場合は特別扱い
            if blk["text"] in ("注", "文献"):
                add_paragraph(doc, blk["text"], font=FONT_BOLD, size_pt=10.0,
                              bold=True, space_before=10, space_after=4)
            else:
                add_paragraph(doc, blk["text"], font=FONT_REGULAR, size_pt=9.0,
                              first_line_indent=True, space_after=2)
        i += 1


def main() -> None:
    if not TEMPLATE.exists():
        raise FileNotFoundError(f"テンプレートが見つかりません: {TEMPLATE}")
    if not DRAFT_MD.exists():
        raise FileNotFoundError(f"原稿が見つかりません: {DRAFT_MD}")

    # テンプレートを開いて中身を空にする（sectPr / styles / theme を保持）
    doc = Document(str(TEMPLATE))
    clear_body_keep_sectPr(doc)

    # 原稿パース
    md_text = DRAFT_MD.read_text(encoding="utf-8")
    blocks = parse_markdown(md_text)

    # レンダリング
    render(blocks, doc)

    # 保存
    doc.save(str(OUT_DOCX))
    print(f"saved: {OUT_DOCX}")
    print(f"blocks rendered: {len(blocks)}")


if __name__ == "__main__":
    main()
